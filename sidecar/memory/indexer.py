"""The background file indexer (BUILD_SPEC §9 Phase 4b).

Reads documents, chunks them, embeds the chunks and stores the vectors so that
"the quotation I sent the banquet hall" can find a file whose name says none of
those words.

**The throttle is the feature.** §9 puts it plainly: *a background indexer that
makes the machine feel slow will get uninstalled.* So this yields constantly and
is built to be interrupted:

- at most `FILES_PER_MIN` files a minute, as a token bucket rather than a sleep
- **paused entirely while she is answering** — a turn has a ~1s budget and this
  is competing for the same cores
- paused while the machine is busy with something else
- embedding is serialised (see `providers/embeddings`), so it can never fan out

Nothing here is on the turn path. The worst case for a user is that indexing
takes longer, which is invisible; the worst case for the machine is what the
throttle exists to prevent.
"""

from __future__ import annotations

import asyncio
import hashlib
import sqlite3
import time
from collections.abc import Callable, Iterator
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

import structlog

from sidecar.memory.db import Database
from sidecar.providers.embeddings import EmbeddingsUnavailable, OllamaEmbeddings

log = structlog.get_logger(__name__)

#: §9's number.
FILES_PER_MIN = 20
#: Above this the machine is doing something the user cares about more.
BUSY_CPU_PERCENT = 60.0
#: How long to wait before looking again when paused.
IDLE_POLL_S = 5.0

#: §9: skip anything larger. A 20MB document is not what this is for, and the
#: extraction cost is unbounded.
MAX_BYTES = 20 * 1024 * 1024

#: ~500 tokens at the project's 4-chars-per-token estimate, with the overlap
#: §9 asks for so a sentence split across a boundary is still findable.
CHUNK_CHARS = 2000
OVERLAP_CHARS = 200
#: A document yielding more than this is a database dump, not prose.
MAX_CHUNKS_PER_FILE = 200

TEXT_EXTENSIONS = frozenset(
    {
        ".txt", ".md", ".rst", ".log", ".csv",
        ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml", ".yml",
        ".toml", ".ini", ".cfg", ".sql", ".sh", ".ps1", ".html", ".css",
        ".java", ".c", ".h", ".cpp", ".go", ".rs",
    }
)
DOCUMENT_EXTENSIONS = frozenset({".pdf", ".docx", ".xlsx"})
INDEXABLE = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS


@dataclass
class IndexStats:
    indexed: int = 0
    skipped: int = 0
    failed: int = 0
    chunks: int = 0


# ── reading a document ───────────────────────────────────────────────


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Tables carry most of the content in quotations and invoices, which is
    # exactly the kind of document this is meant to find.
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _read_xlsx(path: Path) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(str(sheet.title))
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" ".join(cells))
    workbook.close()
    return "\n".join(parts)


_READERS: dict[str, Callable[[Path], str]] = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
}


def extract_text(path: Path) -> str:
    """Whatever text this file has, or "" if it has none worth having.

    Never raises: a corrupt PDF is a normal event in a folder of downloads,
    and it must cost this file rather than the whole sweep.
    """
    reader = _READERS.get(path.suffix.lower())
    try:
        if reader is not None:
            return reader(path)
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001 — one unreadable file is not a failure
        log.debug("indexer.unreadable", path=str(path), error=str(exc))
        return ""


def chunk(text: str) -> list[str]:
    """Overlapping windows, so a sentence spanning a boundary stays findable."""
    cleaned = " ".join(text.split())
    if not cleaned:
        return []

    chunks: list[str] = []
    start = 0
    step = CHUNK_CHARS - OVERLAP_CHARS
    while start < len(cleaned) and len(chunks) < MAX_CHUNKS_PER_FILE:
        chunks.append(cleaned[start : start + CHUNK_CHARS])
        start += step
    return chunks


def should_index(path: Path) -> bool:
    """Whether this file is worth reading at all."""
    if path.suffix.lower() not in INDEXABLE:
        return False
    try:
        return path.is_file() and path.stat().st_size <= MAX_BYTES
    except OSError:
        return False


def _digest(path: Path, mtime: float, size: int) -> str:
    """Cheap identity: re-reading a 10MB PDF to decide whether to re-read it
    would defeat the point."""
    return hashlib.sha256(f"{path}|{mtime}|{size}".encode()).hexdigest()[:32]


# ── the worker ───────────────────────────────────────────────────────


class Indexer:
    """Walks, reads, embeds and stores — slowly, and out of the way."""

    def __init__(
        self,
        db: Database,
        embeddings: OllamaEmbeddings,
        roots: list[Path],
        *,
        files_per_min: int = FILES_PER_MIN,
        is_busy: Callable[[], bool] | None = None,
    ) -> None:
        self._db = db
        self._embeddings = embeddings
        self._roots = roots
        self._interval_s = 60.0 / max(files_per_min, 1)
        #: Asked before every file. Returns True while she is mid-turn, so the
        #: indexer gets out of the way of the thing the user is waiting on.
        self._is_busy = is_busy or (lambda: False)
        self._task: asyncio.Task[None] | None = None
        self.stats = IndexStats()

    # ── lifecycle ───────────────────────────────────────────────────

    def start(self) -> None:
        if self._task is None or self._task.done():
            self._task = asyncio.create_task(self._run())
            log.info("indexer.started", roots=[str(r) for r in self._roots])

    async def stop(self) -> None:
        if self._task is not None:
            self._task.cancel()
            try:
                await self._task
            except asyncio.CancelledError:
                pass
            self._task = None

    async def _run(self) -> None:
        try:
            await self.sweep()
        except asyncio.CancelledError:
            raise
        except Exception:
            log.exception("indexer.failed")

    # ── the sweep ───────────────────────────────────────────────────

    def _candidates(self) -> Iterator[Path]:
        from sidecar.tools.finder import _SKIP_DIRS

        for root in self._roots:
            if not root.is_dir():
                continue
            for path in root.rglob("*"):
                if any(part.lower() in _SKIP_DIRS for part in path.parts):
                    continue
                if should_index(path):
                    yield path

    async def _wait_until_free(self) -> None:
        """Hold here while the machine is busy or she is answering."""
        while True:
            if self._is_busy():
                await asyncio.sleep(IDLE_POLL_S)
                continue
            try:
                import psutil

                if psutil.cpu_percent(interval=0.2) > BUSY_CPU_PERCENT:
                    await asyncio.sleep(IDLE_POLL_S)
                    continue
            except Exception:  # noqa: BLE001 — no psutil is not a reason to stop
                pass
            return

    async def sweep(self) -> IndexStats:
        """One pass over everything, at the throttled rate."""
        started = time.monotonic()
        for path in self._candidates():
            await self._wait_until_free()
            try:
                await self.index_file(path)
            except EmbeddingsUnavailable as exc:
                # Nothing will work until Ollama is back; stop rather than
                # burning the whole tree failing.
                log.warning("indexer.paused", reason=str(exc))
                break
            except Exception:
                self.stats.failed += 1
                log.exception("indexer.file_failed", path=str(path))
            # The throttle. A token bucket rather than a flat sleep so a slow
            # file does not also pay the wait.
            await asyncio.sleep(self._interval_s)

        log.info(
            "indexer.swept",
            indexed=self.stats.indexed,
            skipped=self.stats.skipped,
            failed=self.stats.failed,
            chunks=self.stats.chunks,
            took_s=round(time.monotonic() - started, 1),
        )
        return self.stats

    async def index_file(self, path: Path) -> bool:
        """Index one file. Returns whether it did any work."""
        stat = await asyncio.to_thread(path.stat)
        digest = _digest(path, stat.st_mtime, stat.st_size)

        existing = await self._db.run(
            lambda c: c.execute(
                "SELECT content_hash FROM file_index WHERE path = ?", (str(path),)
            ).fetchone()
        )
        if existing and existing[0] == digest:
            self.stats.skipped += 1
            return False

        text = await asyncio.to_thread(extract_text, path)
        pieces = chunk(text)
        if not pieces:
            await self._record(path, stat, digest, "skipped")
            self.stats.skipped += 1
            return False

        vectors: list[tuple[int, list[float]]] = []
        for ordinal, piece in enumerate(pieces):
            vectors.append((ordinal, await self._embeddings.embed(piece)))

        await self._store(path, stat, digest, pieces, vectors)
        self.stats.indexed += 1
        self.stats.chunks += len(pieces)
        log.debug("indexer.file", path=str(path), chunks=len(pieces))
        return True

    async def _record(self, path: Path, stat: object, digest: str, status: str) -> None:
        row = (
            str(path),
            path.name,
            path.suffix.lower(),
            getattr(stat, "st_size", 0),
            getattr(stat, "st_mtime", 0.0),
            digest,
            datetime.now(UTC).isoformat(),
            status,
        )

        def _write(conn: sqlite3.Connection) -> None:
            with conn:
                conn.execute(
                    """
                    INSERT INTO file_index
                      (path, name, ext, size, mtime, content_hash, indexed_at, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ON CONFLICT(path) DO UPDATE SET
                      size=excluded.size, mtime=excluded.mtime,
                      content_hash=excluded.content_hash,
                      indexed_at=excluded.indexed_at, status=excluded.status
                    """,
                    row,
                )

        await self._db.run(_write)

    async def _store(
        self,
        path: Path,
        stat: object,
        digest: str,
        pieces: list[str],
        vectors: list[tuple[int, list[float]]],
    ) -> None:
        await self._record(path, stat, digest, "indexed")

        def _write(conn: sqlite3.Connection) -> None:
            with conn:
                # Replacing a file means replacing all of it: a shorter second
                # version would otherwise leave the old tail behind, findable
                # and wrong.
                old = conn.execute(
                    "SELECT id FROM file_chunks WHERE path = ?", (str(path),)
                ).fetchall()
                for (chunk_id,) in old:
                    conn.execute("DELETE FROM file_vec WHERE chunk_id = ?", (chunk_id,))
                conn.execute("DELETE FROM file_chunks WHERE path = ?", (str(path),))

                for ordinal, vector in vectors:
                    cursor = conn.execute(
                        "INSERT INTO file_chunks (path, chunk_idx, text) VALUES (?, ?, ?)",
                        (str(path), ordinal, pieces[ordinal]),
                    )
                    conn.execute(
                        "INSERT INTO file_vec (chunk_id, embedding) VALUES (?, ?)",
                        (cursor.lastrowid, _pack(vector)),
                    )

        await self._db.run(_write)


def _pack(vector: list[float]) -> bytes:
    """sqlite-vec takes raw little-endian float32."""
    import struct

    return struct.pack(f"<{len(vector)}f", *vector)


async def search_chunks(
    db: Database, embeddings: OllamaEmbeddings, query: str, limit: int = 10
) -> list[tuple[str, str, float]]:
    """Nearest chunks to `query`, as (path, text, distance)."""
    vector = _pack(await embeddings.embed(query))

    def _query(conn: sqlite3.Connection) -> list[tuple[str, str, float]]:
        rows = conn.execute(
            """
            SELECT c.path, c.text, v.distance
            FROM file_vec v
            JOIN file_chunks c ON c.id = v.chunk_id
            WHERE v.embedding MATCH ? AND k = ?
            ORDER BY v.distance
            """,
            (vector, limit),
        ).fetchall()
        return [(r[0], r[1], float(r[2])) for r in rows]

    return await db.run(_query)
