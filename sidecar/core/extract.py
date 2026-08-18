"""Getting text out of whatever the user hands over.

Eyaas: *"it should be able to handle all office apps extensions, like word,
powerpoint, analyze excel, others... also it should able to analyze
zipped/compressed ones.. and other standard ones too, like literally most of
the widest available ones."*

**Zero new dependencies, and that is not a compromise — it is the same trade
this project has made every time.** `webrtcvad` was declined because
faster-whisper already ships Silero; `beautifulsoup4` was declined and an
HTML parser hand-rolled in `providers/search.py`; `watchdog`, `send2trash`,
`APScheduler` and `pywinauto` all went the same way. The formats added here
are *easier* than HTML:

- **OOXML** (`.pptx`) and **OpenDocument** (`.odt`/`.ods`/`.odp`) and
  **`.epub`** are all zip archives of well-formed XML. `zipfile` +
  `ElementTree`, both standard library.
- **Archives** (`.zip`, `.tar`, `.tar.gz`) are read with `zipfile`/`tarfile`
  and their members handed back to this same registry.
- **`.rtf`** is plain text with control words. A small stripper, the same
  shape and for the same reason as the HTML one.

It also matters that all of it is pure Python: the PyInstaller bundle already
has one unsolved native-library problem (`ctranslate2`), and a parser that
pulls a compiled wheel is a second one waiting.

**What is deliberately NOT supported: `.doc`, `.ppt`, `.xls`** — the pre-2007
OLE2 compound binaries. Every option is bad (`textract` is heavy and
unmaintained, `xlrd` has a CVE history, `antiword` is a binary, COM automation
needs Office installed), and a crude extractor returns exactly the mojibake
that `read_file` was just fixed to stop producing. They are detected by name
and refused **with the fix** — *save it as .docx* — which is worth more to
someone holding a lecture deck than a page of garbled bytes.

**Two extension sets, on purpose.** `INDEXABLE` is what the unattended
background sweep will open; `ATTACHABLE` is what the user may hand over
deliberately. They share every parser and differ in policy: `should_index`
gates a throttled walk over Documents, Desktop and Downloads, and putting
archives in it would have ARIA quietly unpacking every zip on the machine.
"""

from __future__ import annotations

import io
import tarfile
import zipfile
from collections.abc import Callable, Iterator
from pathlib import Path

import structlog

log = structlog.get_logger(__name__)

#: Plain text, read directly. Source files are here because "what did that
#: script do" is a question people ask about their own machine.
TEXT_EXTENSIONS = frozenset(
    {
        ".txt", ".md", ".rst", ".log", ".csv", ".tsv",
        ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml", ".yml",
        ".toml", ".ini", ".cfg", ".sql", ".sh", ".ps1", ".bat",
        ".html", ".htm", ".css", ".java", ".c", ".h", ".cpp", ".go", ".rs",
    }
)

#: Parsed by a reader below. Everything here is either a real dependency the
#: project already carries (pypdf, python-docx, openpyxl) or stdlib.
DOCUMENT_EXTENSIONS = frozenset(
    {".pdf", ".docx", ".xlsx", ".pptx", ".odt", ".ods", ".odp", ".epub", ".rtf"}
)

#: Unpacked and their members read. **Attachments only** — see the module
#: docstring on why the background sweep must not follow these.
ARCHIVE_EXTENSIONS = frozenset({".zip", ".tar", ".gz", ".tgz", ".bz2", ".xz"})

#: The pre-2007 binaries. Named so the refusal can say what to do instead of
#: "unsupported file type".
LEGACY_OFFICE = {
    ".doc": ("Word", ".docx"),
    ".ppt": ("PowerPoint", ".pptx"),
    ".xls": ("Excel", ".xlsx"),
}

#: What the throttled background indexer will open, unchanged from before
#: `.pptx` and friends arrived — plus the ones that cost no more than a PDF.
INDEXABLE = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS

#: What a person may attach. Archives are the only difference, and the
#: difference is the whole point of having two sets.
ATTACHABLE = INDEXABLE | ARCHIVE_EXTENSIONS

# ── archive safety ────────────────────────────────────────────────────
#
# This is the one path in the codebase that unpacks untrusted input, and a zip
# bomb is a plausible thing to be sent rather than a thought experiment. Three
# independent limits, because any one of them alone has a hole: a few enormous
# members, a great many tiny ones, or a nested archive that repeats the
# problem one level down.

#: Members read from a single archive.
MAX_ARCHIVE_MEMBERS = 60
#: Uncompressed bytes read from a single archive, across all members.
MAX_ARCHIVE_BYTES = 25 * 1024 * 1024
#: Archives inside archives are listed, never opened. One level is a mail
#: attachment; two is someone testing what happens.
MAX_ARCHIVE_DEPTH = 1


class Unsupported(Exception):
    """This file cannot be read, and the message says what would work."""


# ── the readers ───────────────────────────────────────────────────────


def _read_pdf(data: bytes, name: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(data: bytes, name: str) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Tables carry most of the content in quotations and invoices, which is
    # exactly the kind of document this is meant to find.
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _read_xlsx(data: bytes, name: str) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(str(sheet.title))
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" ".join(cells))
    workbook.close()
    return "\n".join(parts)


#: OOXML puts every run of text in `<a:t>` under this namespace.
_DRAWING_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"


def _slide_number(name: str) -> int:
    """`ppt/slides/slide10.xml` -> 10.

    **Numeric, not lexical.** Sorting the names as strings gives slide1,
    slide10, slide11, slide2 — a lecture read back in that order looks like
    the model hallucinating rather than like a parser bug, which is the worst
    kind of wrong to debug.
    """
    digits = "".join(c for c in Path(name).stem if c.isdigit())
    return int(digits) if digits else 0


def _read_pptx(data: bytes, name: str) -> str:
    """Slide text and speaker notes, straight out of the OOXML.

    `python-pptx` would do this too, and would pull `lxml` and `XlsxWriter`
    to parse a format the spec guarantees is well-formed XML. That is a worse
    trade than the one this project already accepted when it hand-rolled an
    HTML parser rather than take beautifulsoup4 — HTML is the format with no
    guarantees, and it got the hand-rolled treatment.

    Notes are included because on a lecture deck they are often where the
    actual content is; the slide itself is five words and a diagram.
    """
    from xml.etree import ElementTree

    out: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        slides = sorted(
            (n for n in archive.namelist() if n.startswith("ppt/slides/slide")),
            key=_slide_number,
        )
        notes = sorted(
            (n for n in archive.namelist() if n.startswith("ppt/notesSlides/notesSlide")),
            key=_slide_number,
        )
        for index, member in enumerate(slides, start=1):
            runs = [
                (node.text or "").strip()
                for node in ElementTree.fromstring(archive.read(member)).iter(_DRAWING_NS)
            ]
            body = " ".join(r for r in runs if r)
            if body:
                out.append(f"Slide {index}: {body}")
        for member in notes:
            runs = [
                (node.text or "").strip()
                for node in ElementTree.fromstring(archive.read(member)).iter(_DRAWING_NS)
            ]
            body = " ".join(r for r in runs if r)
            if body:
                out.append(f"Notes: {body}")
    return "\n".join(out)


def _read_opendocument(data: bytes, name: str) -> str:
    """`.odt`, `.ods`, `.odp` — LibreOffice's format, zip + XML like OOXML.

    Every text node in `content.xml` is taken rather than any particular
    element, because the three variants nest text differently and the goal is
    words for a model to read, not a faithful document tree.
    """
    from xml.etree import ElementTree

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        if "content.xml" not in archive.namelist():
            return ""
        root = ElementTree.fromstring(archive.read("content.xml"))
    return "\n".join(t.strip() for t in root.itertext() if t and t.strip())


def _read_epub(data: bytes, name: str) -> str:
    """An epub is a zip of XHTML. Tags are stripped rather than parsed — the
    same call `providers/search.py` already makes about web pages."""
    from sidecar.providers.search import to_text

    out: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        pages = sorted(
            n for n in archive.namelist() if n.lower().endswith((".xhtml", ".html", ".htm"))
        )
        for member in pages[:MAX_ARCHIVE_MEMBERS]:
            out.append(to_text(archive.read(member).decode("utf-8", errors="ignore")))
    return "\n".join(p for p in out if p.strip())


def _read_rtf(data: bytes, name: str) -> str:
    """RTF is text with control words, so this strips rather than parses.

    Not a general RTF implementation and does not claim to be — it drops
    control words, groups and the font/colour tables, and keeps the prose.
    The failure mode is a stray control word in the output, which a model
    handles; the alternative was a dependency for a format that is already
    almost text.
    """
    text = data.decode("utf-8", errors="ignore")
    out: list[str] = []
    depth = 0
    skip_group = -1
    index = 0
    while index < len(text):
        char = text[index]
        if char == "\\":
            # A control word: backslash, letters, optional number.
            index += 1
            word = ""
            while index < len(text) and text[index].isalpha():
                word += text[index]
                index += 1
            while index < len(text) and (text[index].isdigit() or text[index] == "-"):
                index += 1
            if index < len(text) and text[index] == " ":
                index += 1
            if word in ("par", "line", "cell", "row"):
                out.append("\n")
            elif word in ("fonttbl", "colortbl", "stylesheet", "info", "pict"):
                skip_group = depth
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            if skip_group >= 0 and depth <= skip_group + 1:
                skip_group = -1
            depth -= 1
        elif skip_group < 0:
            out.append(char)
        index += 1
    return "".join(out)


_READERS: dict[str, Callable[[bytes, str], str]] = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
    ".pptx": _read_pptx,
    ".odt": _read_opendocument,
    ".ods": _read_opendocument,
    ".odp": _read_opendocument,
    ".epub": _read_epub,
    ".rtf": _read_rtf,
}


# ── archives ──────────────────────────────────────────────────────────


def _safe_member(name: str) -> bool:
    """Whether a member name is one we are willing to read.

    Path traversal is rejected outright. Nothing here writes to disk, so a
    `../../` member cannot escape anywhere — but a name shaped like that says
    something about the archive's intent, and reading it into a prompt is not
    obviously safer than writing it.
    """
    clean = Path(name)
    if clean.is_absolute() or ".." in clean.parts:
        log.warning("extract.archive_traversal", member=name)
        return False
    return True


def _zip_members(data: bytes) -> Iterator[tuple[str, bytes]]:
    seen = 0
    total = 0
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        for info in archive.infolist():
            if seen >= MAX_ARCHIVE_MEMBERS or total >= MAX_ARCHIVE_BYTES:
                return
            if info.is_dir() or not _safe_member(info.filename):
                continue
            with archive.open(info) as handle:
                payload = handle.read(MAX_ARCHIVE_BYTES - total)
            seen += 1
            total += len(payload)
            yield info.filename, payload


def _tar_members(data: bytes) -> Iterator[tuple[str, bytes]]:
    seen = 0
    total = 0
    # `tarfile.open` sniffs the compression, so .tar/.tar.gz/.tgz/.bz2/.xz
    # are all one branch rather than four.
    with tarfile.open(fileobj=io.BytesIO(data)) as archive:
        for info in archive.getmembers():
            if seen >= MAX_ARCHIVE_MEMBERS or total >= MAX_ARCHIVE_BYTES:
                return
            if not info.isfile() or not _safe_member(info.name):
                continue
            handle = archive.extractfile(info)
            if handle is None:
                continue
            with handle:
                payload = handle.read(MAX_ARCHIVE_BYTES - total)
            seen += 1
            total += len(payload)
            yield info.name, payload


def _members(data: bytes, suffix: str) -> Iterator[tuple[str, bytes]]:
    """Every readable member of an archive, within the safety limits."""
    if suffix == ".zip":
        yield from _zip_members(data)
        return
    yield from _tar_members(data)


def _read_archive(data: bytes, name: str, *, depth: int) -> str:
    out: list[str] = []
    listed: list[str] = []
    for member, payload in _members(data, Path(name).suffix.lower()):
        listed.append(member)
        suffix = Path(member).suffix.lower()
        if suffix in ARCHIVE_EXTENSIONS:
            # Listed, not opened. One level is a mail attachment; two is
            # someone finding out what happens.
            continue
        try:
            text = _extract_bytes(payload, member, depth=depth + 1)
        except Unsupported:
            continue
        if text.strip():
            out.append(f"--- {member} ---\n{text}")
    if not out:
        # Still useful: "what is in this zip" is a real question, and an
        # archive of images or binaries has a real answer.
        return "Archive contents:\n" + "\n".join(listed)
    return "\n\n".join(out)


# ── the entry points ──────────────────────────────────────────────────


def _extract_bytes(data: bytes, name: str, *, depth: int = 0) -> str:
    suffix = Path(name).suffix.lower()

    if suffix in LEGACY_OFFICE:
        app, modern = LEGACY_OFFICE[suffix]
        raise Unsupported(
            f"{suffix} is the old binary {app} format, which I cannot read. "
            f"Open it and save as {modern} (File > Save As), or export to PDF, "
            f"and I will read it."
        )

    if suffix in ARCHIVE_EXTENSIONS:
        if depth >= MAX_ARCHIVE_DEPTH:
            raise Unsupported("That is an archive inside an archive; I only open the first.")
        return _read_archive(data, name, depth=depth)

    reader = _READERS.get(suffix)
    if reader is not None:
        return reader(data, name)

    if suffix in TEXT_EXTENSIONS or not suffix:
        return data.decode("utf-8", errors="ignore")

    raise Unsupported(
        f"I cannot read {suffix} files. Documents, spreadsheets, slides, "
        f"plain text, archives and images all work."
    )


def extract_text(path: Path) -> str:
    """Whatever text this file has, or "" if it has none worth having.

    **Never raises**, which is what the background indexer needs: a corrupt
    PDF is a normal event in a folder of downloads and must cost that file
    rather than the whole sweep. Attachments want the opposite — a reason they
    can show the user — and call `extract_or_raise` instead.
    """
    try:
        return extract_or_raise(path)
    except Exception as exc:  # noqa: BLE001 — one unreadable file is not a failure
        log.debug("extract.unreadable", path=str(path), error=str(exc))
        return ""


def extract_or_raise(path: Path) -> str:
    """Same, but an unsupported type raises `Unsupported` with the fix in it.

    The difference matters: a file the *user chose to hand over* failing
    silently is the bug this whole module exists to close. He attached a
    lecture `.ppt`, it was skipped, and the only record was a log line.
    """
    return _extract_bytes(path.read_bytes(), path.name)
